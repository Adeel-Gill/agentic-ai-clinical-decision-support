#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Build a Superior University (DOPS) MS/M.Phil-compliant .docx containing the
front matter + Chapter 1 + Chapter 2 + IEEE references.

Consumes:
  07_Thesis/Compiled/Chapter_1.md
  07_Thesis/Compiled/Chapter_2.md
Produces:
  07_Thesis/Thesis_Ch1_Ch2.docx

Template rules implemented:
  - A4 page, 1" margins except left = 1.25"
  - Times New Roman; body 12pt, 1.5 line spacing, justified
  - Chapter titles: 16pt bold, centered, ALL CAPS
  - Section headings (x.y / x.y.z): 12pt bold, left aligned
  - Front matter in UPPER ROMAN page numbers (Author's Declaration = I)
  - Abstract onward in decimal page numbers, restarting at 1
  - IEEE citation style: [key] -> [n] numbered by order of first appearance;
    reference list emitted in citation order (only cited works appear)
Fill-in placeholders are wrapped in [square brackets] on the title page /
certificates for the student to complete.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
CH1 = os.path.join(HERE, "Chapter_1.md")
CH2 = os.path.join(HERE, "Chapter_2.md")
OUT = os.path.join(ROOT, "07_Thesis", "Thesis_Ch1_Ch2.docx")

# --------------------------------------------------------------------------
# STUDENT-FILLABLE CONFIG  (edit these, then re-run, or edit in Word)
# --------------------------------------------------------------------------
TITLE       = "An Agentic AI Framework for Intelligent Patient Monitoring and Clinical Decision Support"
DEGREE_FULL = "MASTER OF SCIENCE"
SUBJECT     = "ARTIFICIAL INTELLIGENCE"      # <-- verify with your program (e.g., COMPUTER SCIENCE)
STUDENT     = "Adeel Gill"
ROLL_NUMBER = "SU92-MSAIW-S25-011"
SESSION     = "[Session e.g., 2023-2025]"
SUPERVISOR  = "Dr. Fawad Nasim"
DEPARTMENT  = "DEPARTMENT OF COMPUTER SCIENCE"
FACULTY     = "FACULTY OF COMPUTER SCIENCE AND INFORMATION TECHNOLOGY"
UNIVERSITY  = "THE SUPERIOR UNIVERSITY, LAHORE"

FONT = "Times New Roman"

# --------------------------------------------------------------------------
# IEEE reference strings keyed by bibtex key (derived from References.bib).
# Only keys actually cited in the chapters are emitted, numbered by appearance.
# VERIFY details against the publisher before final submission.
# --------------------------------------------------------------------------
IEEE = {
 "yao2023react": 'S. Yao, J. Zhao, D. Yu, N. Du, I. Shafran, K. Narasimhan, and Y. Cao, "ReAct: Synergizing reasoning and acting in language models," in Proc. Int. Conf. Learn. Represent. (ICLR), 2023.',
 "park2023generative": 'J. S. Park, J. C. O’Brien, C. J. Cai, M. R. Morris, P. Liang, and M. S. Bernstein, "Generative agents: Interactive simulacra of human behavior," in Proc. 36th ACM Symp. User Interface Softw. Technol. (UIST), 2023.',
 "xi2023rise": 'Z. Xi et al., "The rise and potential of large language model based agents: A survey," arXiv:2309.07864, 2023.',
 "schick2023toolformer": 'T. Schick et al., "Toolformer: Language models can teach themselves to use tools," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
 "wang2023voyager": 'G. Wang et al., "Voyager: An open-ended embodied agent with large language models," Trans. Mach. Learn. Res. (TMLR), 2024.',
 "wu2024autogen": 'Q. Wu et al., "AutoGen: Enabling next-gen LLM applications via multi-agent conversation," in Proc. Conf. Lang. Model. (COLM), 2024.',
 "li2023camel": 'G. Li, H. A. A. K. Hammoud, H. Itani, D. Khizbullin, and B. Ghanem, "CAMEL: Communicative agents for ‘mind’ exploration of large language model society," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
 "hong2024metagpt": 'S. Hong et al., "MetaGPT: Meta programming for a multi-agent collaborative framework," in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.',
 "wang2024survey": 'L. Wang et al., "A survey on large language model based autonomous agents," Front. Comput. Sci., vol. 18, no. 6, p. 186345, 2024.',
 "wei2022chain": 'J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2022.',
 "wang2023selfconsistency": 'X. Wang et al., "Self-consistency improves chain of thought reasoning in language models," in Proc. Int. Conf. Learn. Represent. (ICLR), 2023.',
 "yao2023tree": 'S. Yao et al., "Tree of thoughts: Deliberate problem solving with large language models," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
 "shinn2023reflexion": 'N. Shinn, F. Cassano, E. Berman, A. Gopinath, K. Narasimhan, and S. Yao, "Reflexion: Language agents with verbal reinforcement learning," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
 "sapkota2025agents": 'R. Sapkota, K. I. Roumeliotis, and M. Karkee, "AI agents vs. Agentic AI: A conceptual taxonomy, applications and challenges," arXiv:2505.10468, 2025.',
 "lewis2020rag": 'P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2020.',
 "gao2023rag": 'Y. Gao et al., "Retrieval-augmented generation for large language models: A survey," arXiv:2312.10997, 2023.',
 "asai2024selfrag": 'A. Asai, Z. Wu, Y. Wang, A. Sil, and H. Hajishirzi, "Self-RAG: Learning to retrieve, generate, and critique through self-reflection," in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.',
 "zhao2025medrag": 'X. Zhao, S. Liu, S.-Y. Yang, and C. Miao, "MedRAG: Enhancing retrieval-augmented generation with knowledge graph-elicited reasoning for healthcare copilot," in Proc. ACM Web Conf. (WWW), 2025.',
 "singhal2023clinical": 'K. Singhal et al., "Large language models encode clinical knowledge," Nature, vol. 620, no. 7972, pp. 172–180, 2023.',
 "singhal2025medpalm2": 'K. Singhal et al., "Toward expert-level medical question answering with large language models," Nat. Med., vol. 31, pp. 943–950, 2025.',
 "tu2024generalist": 'T. Tu et al., "Towards generalist biomedical AI," NEJM AI, vol. 1, no. 3, 2024.',
 "thirunavukarasu2023llms": 'A. J. Thirunavukarasu, D. S. J. Ting, K. Elangovan, L. Gutierrez, T. F. Tan, and D. S. W. Ting, "Large language models in medicine," Nat. Med., vol. 29, no. 8, pp. 1930–1940, 2023.',
 "zhou2024survey": 'H. Zhou et al., "A survey of large language models in medicine: Progress, application, and challenge," arXiv:2311.05112, 2024.',
 "tang2024medagents": 'X. Tang, A. Zou, Z. Zhang, Z. Li, Y. Zhao, X. Zhang, A. Cohan, and M. Gerstein, "MedAgents: Large language models as collaborators for zero-shot medical reasoning," in Findings Assoc. Comput. Linguist. (ACL Findings), 2024.',
 "li2024agenthospital": 'J. Li et al., "Agent hospital: A simulacrum of hospital with evolvable medical agents," arXiv:2405.02957, 2024.',
 "toma2023clinicalcamel": 'A. Toma, P. R. Lawler, J. Ba, R. G. Krishnan, B. B. Rubin, and B. Wang, "Clinical Camel: An open expert-level medical language model with dialogue-based knowledge encoding," arXiv:2305.12031, 2023.',
 "jin2021medqa": 'D. Jin, E. Pan, N. Oufattole, W.-H. Weng, H. Fang, and P. Szolovits, "What disease does this patient have? A large-scale open domain question answering dataset from medical exams," Appl. Sci., vol. 11, no. 14, p. 6421, 2021.',
 "shi2024ehragent": 'W. Shi et al., "EHRAgent: Code empowers large language models for few-shot complex tabular reasoning on electronic health records," in Proc. EMNLP, 2024.',
 "schmidgall2024agentclinic": 'S. Schmidgall, R. Ziaei, C. Harris, E. Reis, J. Jopling, and M. Moor, "AgentClinic: A multimodal agent benchmark to evaluate AI in simulated clinical environments," arXiv:2405.07960, 2024.',
 "tu2025amie": 'T. Tu et al., "Towards conversational diagnostic AI," Nature, 2025. [venue to verify]',
 "johnson2023mimic": 'A. E. W. Johnson et al., "MIMIC-IV, a freely accessible electronic health record dataset," Sci. Data, vol. 10, no. 1, p. 1, 2023.',
 "goldberger2000physiobank": 'A. L. Goldberger et al., "PhysioBank, PhysioToolkit, and PhysioNet: Components of a new research resource for complex physiologic signals," Circulation, vol. 101, no. 23, pp. e215–e220, 2000.',
 "rasheed2022explainable": 'K. Rasheed, A. Qayyum, M. Ghaly, A. Al-Fuqaha, A. Razi, and J. Qadir, "Explainable, trustworthy, and ethical machine learning for healthcare: A survey," Comput. Biol. Med., vol. 149, p. 106043, 2022.',
 "jimenez2023trustworthy": 'Author(s), "Toward trustworthy AI in healthcare," 2023. [full citation to verify against stored P018 PDF]',
 "hevner2004design": 'A. R. Hevner, S. T. March, J. Park, and S. Ram, "Design science in information systems research," MIS Quart., vol. 28, no. 1, pp. 75–105, 2004.',
 "singer2016sepsis3": 'M. Singer et al., "The third international consensus definitions for sepsis and septic shock (Sepsis-3)," JAMA, vol. 315, no. 8, pp. 801–810, 2016.',
 "vincent1996sofa": 'J.-L. Vincent et al., "The SOFA score to describe organ dysfunction/failure," Intensive Care Med., vol. 22, no. 7, pp. 707–710, 1996.',
 "es2024ragas": 'S. Es, J. James, L. Espinosa-Anke, and S. Schockaert, "RAGAS: Automated evaluation of retrieval augmented generation," in Proc. EACL: System Demonstrations, 2024.',
 "guo2017calibration": 'C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, "On calibration of modern neural networks," in Proc. Int. Conf. Mach. Learn. (ICML), 2017.',
 "hardt2016equality": 'M. Hardt, E. Price, and N. Srebro, "Equality of opportunity in supervised learning," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2016.',
 "mehrabi2021survey": 'N. Mehrabi, F. Morstatter, N. Saxena, K. Lerman, and A. Galstyan, "A survey on bias and fairness in machine learning," ACM Comput. Surv., vol. 54, no. 6, pp. 1–35, 2021.',
 "delong1988comparing": 'E. R. DeLong, D. M. DeLong, and D. L. Clarke-Pearson, "Comparing the areas under two or more correlated ROC curves: A nonparametric approach," Biometrics, vol. 44, no. 3, pp. 837–845, 1988.',
 "dietterich1998approximate": 'T. G. Dietterich, "Approximate statistical tests for comparing supervised classification learning algorithms," Neural Comput., vol. 10, no. 7, pp. 1895–1923, 1998.',
 "benjamini1995controlling": 'Y. Benjamini and Y. Hochberg, "Controlling the false discovery rate: A practical and powerful approach to multiple testing," J. R. Stat. Soc. B, vol. 57, no. 1, pp. 289–300, 1995.',
 "hayes2007krippendorff": 'A. F. Hayes and K. Krippendorff, "Answering the call for a standard reliability measure for coding data," Commun. Methods Meas., vol. 1, no. 1, pp. 77–89, 2007.',
}

ABSTRACT = (
 "Clinical teams in intensive care generate vast, heterogeneous data—vital signs, laboratory "
 "results, medications, and free-text notes—yet turning this stream into timely, trustworthy "
 "decisions remains difficult. Large language models (LLMs) answer medical questions well but "
 "behave as passive responders: they lack persistent patient memory, autonomous planning, and "
 "verifiable grounding, and they are typically evaluated on static examination-style questions "
 "rather than on the longitudinal record of a real patient. This thesis designs an Agentic AI "
 "framework for intelligent patient monitoring and clinical decision support that addresses these "
 "limitations. The framework organizes specialized LLM agents—monitoring, diagnosis, risk "
 "prediction, treatment recommendation, explanation, and a verification gate—under a coordinator "
 "that routes tasks by patient condition and resolves conflicting recommendations. It couples the "
 "ReAct reasoning paradigm with retrieval-augmented generation (RAG) that grounds every "
 "recommendation in both the patient’s own electronic health record timeline and external clinical "
 "guidelines, and it preserves longitudinal context through a layered memory model. A cross-cutting "
 "trustworthy-AI layer provides explainability, audit logging, bias monitoring, confidence "
 "calibration, and human-in-the-loop validation so that clinicians retain final authority. The "
 "publicly available MIMIC-IV critical-care database is used as the data source, and the study "
 "specifies a bounded prototype together with an evaluation protocol that compares the framework "
 "against single-LLM, guideline-only RAG, and multi-agent baselines using measures of diagnostic "
 "accuracy, factual grounding, risk-prediction discrimination, safety, explanation faithfulness, "
 "calibration, and runtime cost. The expected contribution is a coherent, verifiable architecture "
 "for longitudinal clinical decision support, an explicit account of how patient-timeline retrieval "
 "and a dedicated verification agent affect hallucination and safety, and a reproducible design that "
 "future work can implement and validate. The thesis is positioned as a design and methodology "
 "contribution that lays the foundation for empirical clinical evaluation."
)
KEYWORDS = ("Keywords: Agentic AI; large language models; multi-agent systems; retrieval-augmented "
            "generation; clinical decision support; patient monitoring; MIMIC-IV; trustworthy AI.")

ABBREVIATIONS = [
 ("AI", "Artificial Intelligence"),
 ("CDSS", "Clinical Decision Support System"),
 ("CoT", "Chain-of-Thought"),
 ("EHR", "Electronic Health Record"),
 ("HITL", "Human-in-the-Loop"),
 ("ICU", "Intensive Care Unit"),
 ("LLM", "Large Language Model"),
 ("MIMIC", "Medical Information Mart for Intensive Care"),
 ("RAG", "Retrieval-Augmented Generation"),
 ("ReAct", "Reasoning and Acting"),
 ("XAI", "Explainable Artificial Intelligence"),
]

# --------------------------------------------------------------------------
# Low-level helpers
# --------------------------------------------------------------------------
def set_cell_no_border(table):
    tbl = table._tbl
    for el in tbl.iter():
        pass

def _run_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)

def para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False,
         italic=False, spacing=1.5, before=0, after=6, caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if text:
        r = p.add_run(text.upper() if caps else text)
        _run_font(r, size=size, bold=bold, italic=italic)
    return p

def heading_block(doc, label, title):
    """Chapter heading: two centered ALL-CAPS 16pt bold lines."""
    para(doc, label, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
         before=0, after=0, caps=True)
    para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
         before=0, after=12, caps=True)

def mandatory_title(doc, text):
    para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
         before=0, after=18, caps=True)

def new_page(doc):
    doc.add_page_break()

# ---- page numbering (section-level) --------------------------------------
def set_pgnum(section, fmt=None, start=None):
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    el = OxmlElement('w:pgNumType')
    if fmt:
        el.set(qn('w:fmt'), fmt)
    if start is not None:
        el.set(qn('w:start'), str(start))
    sectPr.append(el)

def footer_page_field(section, align=WD_ALIGN_PARAGRAPH.CENTER, show=True):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.text = ""
    p.alignment = align
    if not show:
        return
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = "PAGE"
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(instr); run._r.append(e)
    _run_font(run, size=12)

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Right-click and choose “Update Field” to build the table of contents."
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(instr); run._r.append(sep); run._r.append(t); run._r.append(e)
    _run_font(run, size=12)

# --------------------------------------------------------------------------
# Citation handling + inline runs
# --------------------------------------------------------------------------
CITE_RE = re.compile(r"\[([A-Za-z]+\d{4}[a-z0-9]*)\]")

def substitute_citations(text, cite_map, order):
    def repl(m):
        key = m.group(1)
        if key not in IEEE:
            return m.group(0)  # leave unknown keys untouched (visible flag)
        if key not in cite_map:
            cite_map[key] = len(order) + 1
            order.append(key)
        return "[%d]" % cite_map[key]
    return CITE_RE.sub(repl, text)

def add_inline(p, text):
    """Split on ** for bold; everything else normal 12pt TNR."""
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        r = p.add_run(seg)
        _run_font(r, size=12, bold=(i % 2 == 1))

def body_para(doc, text, cite_map, order, bullet=False):
    text = substitute_citations(text, cite_map, order)
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    add_inline(p, text)
    return p

def section_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    r = p.add_run(text)
    _run_font(r, size=12, bold=True)

def table_placeholder(doc, text):
    p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True, after=6)
    return p

# --------------------------------------------------------------------------
# Chapter markdown parser (contract defined in the compile step)
# --------------------------------------------------------------------------
def render_chapter(doc, path, cite_map, order):
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    para_buf = []
    def flush():
        if para_buf:
            body_para(doc, " ".join(para_buf).strip(), cite_map, order)
            para_buf.clear()

    for raw in lines:
        line = raw.rstrip()
        s = line.strip()
        if s == "":
            flush(); continue
        if s.startswith("```") or s.startswith("|") or set(s) <= set("-= "):
            # skip code fences, stray table rows, horizontal rules
            continue
        if s.startswith("# "):
            flush()
            head = s[2:].strip()
            if ":" in head:
                label, title = head.split(":", 1)
            else:
                label, title = head, ""
            heading_block(doc, label.strip(), title.strip())
            continue
        if s.startswith("### "):
            flush(); section_heading(doc, s[4:].strip()); continue
        if s.startswith("## "):
            flush(); section_heading(doc, s[3:].strip()); continue
        if s.startswith("[TABLE:"):
            flush()
            inner = s[1:-1] if s.endswith("]") else s[1:]
            inner = inner.replace("TABLE:", "").strip()
            table_placeholder(doc, "[" + inner + " — insert table here]")
            continue
        if s.startswith("- ") or s.startswith("* "):
            flush(); body_para(doc, s[2:].strip(), cite_map, order, bullet=True); continue
        para_buf.append(s)
    flush()

# --------------------------------------------------------------------------
# Front matter
# --------------------------------------------------------------------------
def build_title_page(doc):
    para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, before=24, after=24)
    para(doc, "[Insert Superior University logo here — 5 cm × 7.43 cm]",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True, after=24)
    para(doc, "A Thesis submitted in partial fulfillment", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=0)
    para(doc, "of the requirements for the degree of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=12)
    para(doc, DEGREE_FULL, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0, caps=True)
    para(doc, "IN", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0)
    para(doc, SUBJECT, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=18, caps=True)
    para(doc, "Submitted By", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=0)
    para(doc, STUDENT, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0)
    para(doc, ROLL_NUMBER, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0)
    para(doc, "Session: " + SESSION, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=18)
    para(doc, "Supervised By", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=0)
    para(doc, SUPERVISOR, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=24)
    para(doc, DEPARTMENT, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0, caps=True)
    para(doc, FACULTY, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0, caps=True)
    para(doc, UNIVERSITY, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=0, caps=True)

def sig_line(doc, label):
    para(doc, label + " ____________________________", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, after=6)

def build_declaration(doc):
    mandatory_title(doc, "Author's Declaration")
    body = ('I, %s, a student of “%s %s” at the “%s”, The Superior University, Lahore, '
            'hereby declare that this thesis titled, “%s” is my own research work and has not '
            'been submitted, published, or printed elsewhere in Pakistan or abroad. Additionally, I '
            'will not use this thesis to obtain any degree other than the one stated above.'
            % (STUDENT, "MS", SUBJECT.title(), FACULTY.title(), TITLE))
    body_para(doc, body, {}, [])
    body_para(doc, ('I fully understand that if my statement is found to be incorrect at any stage, '
                    'including after the award of the degree, the university has the right to revoke my '
                    'MS/M.Phil degree.'), {}, [])
    for _ in range(2):
        para(doc, "", after=0)
    sig_line(doc, "Signature of Student:")
    sig_line(doc, "Name of Student:      " + STUDENT)
    sig_line(doc, "Roll Number:          " + ROLL_NUMBER)
    sig_line(doc, "Date:")

def build_plagiarism(doc):
    mandatory_title(doc, "Plagiarism Undertaking")
    body_para(doc, ('I solemnly declare that the research work presented in this thesis titled, “%s” '
                    'is solely my research work, and that the entire thesis has been completed by me, with '
                    'no significant contribution from any other person or institution. Any small '
                    'contribution, wherever taken, has been duly acknowledged.' % TITLE), {}, [])
    body_para(doc, ('I understand the zero-tolerance policy of the HEC and Superior University, Lahore, '
                    'towards plagiarism. Therefore, I, as the author of the above-titled thesis, declare '
                    'that no portion of my thesis has been plagiarized and that all material used from '
                    'other sources has been properly acknowledged, cited, and referenced.'), {}, [])
    body_para(doc, ('I undertake that if I am found guilty of any formal plagiarism in the above titled '
                    'thesis, even after the award of the MS/M.Phil degree, the University reserves the '
                    'right to revoke my degree, and that HEC and the University have the right to publish '
                    'my name on the HEC/University website for submitting a plagiarized thesis.'), {}, [])
    para(doc, "", after=0)
    sig_line(doc, "Signature of Student:")
    sig_line(doc, "Name of Student:      " + STUDENT)
    para(doc, "", after=0)
    sig_line(doc, "Signature Supervisor:")
    sig_line(doc, "Supervisor Name:      " + SUPERVISOR)

def build_completion(doc):
    mandatory_title(doc, "Certificate of Research Completion")
    body_para(doc, ('It is certified that this thesis titled, “%s”, submitted by %s, Roll No. %s, '
                    'for the MS/M.Phil “%s” at the “%s”, Superior University, is an original '
                    'research work and contains satisfactory material to be eligible for evaluation by the '
                    'Examiner(s) for the award of the above-stated degree.'
                    % (TITLE, STUDENT, ROLL_NUMBER, SUBJECT.title(), FACULTY.title())), {}, [])
    body_para(doc, "No part of this dissertation has been submitted anywhere else for any other degree.", {}, [])
    para(doc, "", after=0)
    sig_line(doc, "Supervisor's Signature:")
    sig_line(doc, "Supervisor's Name:     " + SUPERVISOR)
    para(doc, "Designation:", size=12, after=6)
    para(doc, "", after=0)
    sig_line(doc, "Co-Supervisor's Signature:")
    sig_line(doc, "Co-Supervisor's Name:")
    para(doc, "Designation:", size=12, after=6)

def build_approval(doc):
    mandatory_title(doc, "Certificate of Approval")
    body_para(doc, ('It is certified that the research work contained in this thesis, titled “%s”, '
                    'was conducted by %s under the supervision of %s and is adequate for the award of '
                    '“%s in %s”.' % (TITLE, STUDENT, SUPERVISOR, DEGREE_FULL.title(), SUBJECT.title())),
              {}, [])
    body_para(doc, ('No part of this thesis has been submitted anywhere else to any other degree. This '
                    'thesis is submitted to the “%s”, The Superior University, Lahore, in partial '
                    'fulfillment of the requirements for the degree.' % FACULTY.title()), {}, [])
    para(doc, "", after=0)
    para(doc, "Student Name: " + STUDENT + "     Signature: ____________________", size=12, after=6)
    para(doc, "Examination Committee:", size=12, bold=True, after=6)
    for role in ["a) Session Chair:", "b) External Examiner (Name / Designation / University):",
                 "c) Internal Examiner (Name / Designation / University):",
                 "d) Supervisor:", "e) Head of Department:", "f) Dean:", "g) COE:"]:
        para(doc, role + "     Signature: ____________________", size=12, after=6)

def build_dedication(doc):
    mandatory_title(doc, "Dedication")
    para(doc, "To [dedication text — begins with the word “To”].",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)

def build_acknowledgements(doc):
    mandatory_title(doc, "Acknowledgements")
    body_para(doc, ("[Acknowledge your supervisor, collaborators, department, family, and any funding "
                    "source here. This page is optional and can be personalized before submission.]"), {}, [])

def build_lists(doc):
    mandatory_title(doc, "List of Figures")
    body_para(doc, ("(Insert figure captions here after placing figures in the document. In Word: "
                    "References → Insert Table of Figures. Example: Fig 3.1: Proposed Agentic AI "
                    "framework.)"), {}, [])
    new_page(doc)
    mandatory_title(doc, "List of Tables")
    body_para(doc, ("(Insert table captions here. Example: Table 2.1: Comparative analysis of reviewed "
                    "agent frameworks.)"), {}, [])
    new_page(doc)
    mandatory_title(doc, "List of Abbreviations and Acronyms")
    tbl = doc.add_table(rows=0, cols=2)
    for ab, full in ABBREVIATIONS:
        cells = tbl.add_row().cells
        r = cells[0].paragraphs[0].add_run(ab); _run_font(r, size=12, bold=True)
        r2 = cells[1].paragraphs[0].add_run(full); _run_font(r2, size=12)

def build_abstract(doc, cite_map, order):
    mandatory_title(doc, "Abstract")
    body_para(doc, ABSTRACT, cite_map, order)
    body_para(doc, KEYWORDS, cite_map, order)

# --------------------------------------------------------------------------
# References
# --------------------------------------------------------------------------
def build_references(doc, order):
    mandatory_title(doc, "References")
    for i, key in enumerate(order, start=1):
        text = IEEE.get(key, "[MISSING REFERENCE: %s]" % key)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(6)
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)  # hanging indent
        r = p.add_run("[%d]  " % i); _run_font(r, size=12)
        r2 = p.add_run(text); _run_font(r2, size=12)

# --------------------------------------------------------------------------
# Assemble
# --------------------------------------------------------------------------
def main():
    doc = Document()
    # default font
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    # ---- Section 0: Title page ----
    build_title_page(doc)

    # ---- Section 1: front matter (roman) ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_declaration(doc);   new_page(doc)
    build_plagiarism(doc);    new_page(doc)
    build_completion(doc);    new_page(doc)
    build_approval(doc);      new_page(doc)
    build_dedication(doc);    new_page(doc)
    build_acknowledgements(doc); new_page(doc)
    mandatory_title(doc, "Table of Contents"); add_toc(doc); new_page(doc)
    build_lists(doc)

    # ---- Section 2: body (decimal, restart at 1) ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    cite_map, order = {}, []
    build_abstract(doc, cite_map, order); new_page(doc)
    render_chapter(doc, CH1, cite_map, order); new_page(doc)
    render_chapter(doc, CH2, cite_map, order); new_page(doc)
    build_references(doc, order)

    # ---- page setup + numbering for all sections ----
    secs = doc.sections
    for s in secs:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.right_margin = Inches(1)
        s.left_margin = Inches(1.25)
    # title page: no number
    footer_page_field(secs[0], show=False)
    set_pgnum(secs[0], fmt="upperRoman", start=1)
    # front matter: upper roman starting I
    footer_page_field(secs[1], align=WD_ALIGN_PARAGRAPH.RIGHT, show=True)
    set_pgnum(secs[1], fmt="upperRoman", start=1)
    # body: decimal restart at 1
    footer_page_field(secs[2], align=WD_ALIGN_PARAGRAPH.RIGHT, show=True)
    set_pgnum(secs[2], fmt="decimal", start=1)

    doc.save(OUT)
    print("Saved:", OUT)
    print("References cited (in order):", len(order))
    for i, k in enumerate(order, 1):
        print("  [%d] %s" % (i, k))

if __name__ == "__main__":
    main()
