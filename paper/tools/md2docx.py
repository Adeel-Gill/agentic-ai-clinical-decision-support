"""Build the W-paper as a .docx from the markdown sections.

Citations [key; key2] are numbered in order of first appearance (matching the
IEEEtran PDF) and rendered as [n]; a References section is generated from
references.bib. Run from the repository root:

    python paper/tools/md2docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ORDER = ["01_Introduction.md", "02_Related_Work.md", "03_Methodology.md",
         "04_Discussion.md", "05_Pilot_Study.md", "06_Conclusion.md"]
PAPER = Path("paper")
TITLE = ("An Agentic AI Framework for Intelligent Patient Monitoring and "
         "Clinical Decision Support with Patient-Timeline Retrieval and "
         "Verified Recommendations")
AUTHOR = ("Adeel Masih — Department of Computer Science and IT, The Superior University, "
          "Lahore, Pakistan · Supervisor: Dr. Fawad Nasim · adeel_gill@outlook.com")

CITE_RE = re.compile(r"\[([a-z][a-zA-Z0-9]+(?:;\s*[a-z][a-zA-Z0-9]+)*)\]")


# ---------------- citation numbering ----------------
def collect_citation_order() -> dict[str, int]:
    order: dict[str, int] = {}
    for f in ORDER:
        for m in CITE_RE.finditer((PAPER / f).read_text(encoding="utf-8")):
            for k in re.split(r";\s*", m.group(1)):
                if k not in order:
                    order[k] = len(order) + 1
    return order


# ---------------- bib parsing ----------------
LATEX_ACCENTS = [
    (r"\{\\c\{c\}\}", "ç"), (r"\\c\{c\}", "ç"),
    (r"\{\\'e\}", "é"), (r"\\'e", "é"), (r"\{\\'a\}", "á"),
    (r'\{\\"o\}', "ö"), (r'\{\\"u\}', "ü"), (r"\{\\~n\}", "ñ"),
]


def parse_bib() -> dict[str, dict]:
    txt = (PAPER / "references.bib").read_text(encoding="utf-8")
    # normalize LaTeX accent macros to unicode BEFORE brace-based field parsing,
    # otherwise nested groups like {\c{c}} truncate the field match
    for pat, rep in LATEX_ACCENTS:
        txt = re.sub(pat, rep, txt)
    entries = {}
    # field regex tolerates two levels of brace nesting
    field_re = re.compile(r"(\w+)\s*=\s*\{((?:[^{}]|\{(?:[^{}]|\{[^{}]*\})*\})*)\}")
    for raw in re.split(r"(?=^@)", txt, flags=re.M):
        m = re.match(r"^@(\w+)\{([^,]+),", raw)
        if not m:
            continue
        fields = dict(field_re.findall(raw))
        fields = {k.lower(): re.sub(r"[{}]", "", v).strip() for k, v in fields.items()}
        entries[m.group(2)] = {"type": m.group(1), **fields}
    return entries


def format_ref(e: dict) -> str:
    authors = e.get("author", "")
    parts_a = []
    for a in authors.split(" and "):
        a = a.strip()
        if a == "others":
            parts_a.append("et al.")
            continue
        if "," in a:
            last, first = [x.strip() for x in a.split(",", 1)]
            initials = " ".join(w[0] + "." for w in first.split() if w)
            parts_a.append(f"{initials} {last}".strip())
        else:
            parts_a.append(a)
    astr = ", ".join(parts_a)
    title = e.get("title", "")
    venue = e.get("journal") or e.get("booktitle") or ""
    bits = [f'{astr}, "{title}," {venue}'.rstrip(", ")]
    for key, prefix in [("volume", "vol. "), ("number", "no. "), ("pages", "pp. ")]:
        if e.get(key):
            bits.append(prefix + e[key])
    if e.get("year"):
        bits.append(e["year"])
    ref = ", ".join(bits)
    if e.get("doi"):
        ref += f", doi: {e['doi']}"
    return ref + "."


# ---------------- inline rendering ----------------
def add_runs(par, text: str, cites: dict[str, int]) -> None:
    def cite_sub(m):
        nums = sorted(cites[k] for k in re.split(r";\s*", m.group(1)))
        return "[" + "], [".join(str(n) for n in nums) + "]"

    text = CITE_RE.sub(cite_sub, text)
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text, flags=re.S):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        par.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def main() -> None:
    cites = collect_citation_order()
    bib = parse_bib()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(15)
    a = doc.add_paragraph(AUTHOR)
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # abstract + keywords
    abs_md = (PAPER / "00_Abstract.md").read_text(encoding="utf-8")
    abstract = re.search(r"# Abstract\n\n(.*?)\n\n# Keywords", abs_md, re.S).group(1)
    keywords = abs_md.split("# Keywords")[1].strip()
    h = doc.add_paragraph()
    h.add_run("Abstract— ").bold = True
    add_runs(h, " ".join(abstract.split()), cites)
    k = doc.add_paragraph()
    k.add_run("Index Terms— ").bold = True
    k.add_run(" ".join(keywords.split()))

    sec_no = 0
    for f in ORDER:
        text = (PAPER / f).read_text(encoding="utf-8")
        sub_no = 0
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("Table: "):
                lines = block.split("\n")
                cap = doc.add_paragraph()
                cr = cap.add_run("TABLE. " + lines[0][len("Table: "):])
                cr.bold = True
                cr.font.size = Pt(9)
                rows = [[c.strip() for c in l.strip().strip("|").split("|")]
                        for l in lines[1:] if l.startswith("|")]
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        c = tbl.cell(ri, ci)
                        c.text = cell
                        for par in c.paragraphs:
                            for run in par.runs:
                                run.font.size = Pt(9)
                                if ri == 0:
                                    run.bold = True
                continue
            m1 = re.match(r"^# \d+\.\s*(.+)$", block)
            m2 = re.match(r"^## \d+\.\d+\s*(.+)$", block)
            if m1:
                sec_no += 1
                sub_no = 0
                doc.add_heading(f"{sec_no}. {m1.group(1)}", level=1)
            elif m2:
                sub_no += 1
                doc.add_heading(f"{sec_no}.{sub_no} {m2.group(1)}", level=2)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_runs(p, " ".join(block.split()), cites)
            # figure after the Design Overview subsection
            if m2 and m2.group(1).strip() == "Design Overview":
                pass  # figure inserted after the following paragraph
        if f == "03_Methodology.md":
            pass

    # insert figure after first Methodology paragraph is complex positionally;
    # append it at the end of the document body instead, before references.
    doc.add_heading("Figure 1", level=2)
    doc.add_picture(str(PAPER / "fig_architecture.png"), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(
        "Fig. 1. The proposed agentic AI framework over MIMIC-IV: six layers closed by "
        "human-in-the-loop validation, governed by a cross-cutting trustworthy-AI panel.")
    cr.font.size = Pt(9)

    doc.add_heading("Ethics Statement", level=1)
    doc.add_paragraph(
        "This study uses only the MIMIC-IV Clinical Database Demo v2.2, a fully de-identified, "
        "openly licensed public dataset distributed by PhysioNet that requires no credentialing "
        "and involves no interaction with patients or clinicians; no additional ethics approval "
        "was required for its use. The planned full evaluation will use the credentialed MIMIC-IV "
        "database under the PhysioNet Credentialed Health Data Use Agreement and completed human "
        "subjects research training. The system described is an academic proof-of-concept "
        "evaluated on de-identified and synthetic data only; it is not a medical device and is "
        "not intended for clinical use, and any output it produces requires review by a "
        "qualified licensed clinician.")
    doc.add_heading("Data Availability", level=1)
    doc.add_paragraph(
        "The MIMIC-IV Clinical Database Demo v2.2 is publicly available from PhysioNet. The "
        "pilot implementation (timeline construction, timestamp-aware retrieval, early-warning "
        "baseline, and verification gate) and the aggregate pilot metrics reported in Section 5 "
        "are available in the project repository; no patient-level data is redistributed. The "
        "full MIMIC-IV database is available to credentialed researchers via PhysioNet.")

    doc.add_heading("Funding", level=1)
    doc.add_paragraph("The authors received no specific funding for this work.")
    doc.add_heading("Competing Interests", level=1)
    doc.add_paragraph("The authors declare no competing interests.")
    doc.add_heading("Author Contributions", level=1)
    doc.add_paragraph(
        "A.M. designed the framework, implemented the pilot study and platform prototype, and "
        "wrote the manuscript. F.N. supervised the research and reviewed the manuscript.")
    doc.add_heading("AI-Assistance Disclosure", level=1)
    doc.add_paragraph(
        "Large language model tools were used, under the authors' direction, to assist with "
        "literature search, drafting, and editing of this manuscript and with the "
        "implementation of the pilot software. The authors reviewed and verified all content, "
        "citations, analyses, and results, and take full responsibility for the manuscript.")

    doc.add_heading("References", level=1)
    for key, n in sorted(cites.items(), key=lambda kv: kv[1]):
        e = bib.get(key)
        p = doc.add_paragraph(f"[{n}] " + (format_ref(e) if e else key))
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(10)

    out = PAPER / "Agentic_AI_Framework_for_Intelligent_Patient_Monitoring_and_Clinical_Decision_Support.docx"
    doc.save(out)
    print("written:", out, "| citations:", len(cites))


if __name__ == "__main__":
    main()
